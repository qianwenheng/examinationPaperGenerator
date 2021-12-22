import statistics.subject
import os
import random
import docx
import utils.docAdding

questionChineseNo = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、', '十一、', '十二、', '十三、', '十四、', '十五、',
                     '十六、', '十七、', '十八、', '十九、', '二十、']


class PaperInfo:
    def __init__(self, subjectInfo=None):
        self.subjectInfo = subjectInfo
        self.questionTypes = []  # 考卷包含的所有题型,ie ['选择题','填空题']
        self.questionNums = []  # i.e. [20,20]

        self.chapterQuestionNums = []
        self.questionChapterNums = []
        self.questionTypeQuestionNames = []
        self.chapterQuestionNames = []
        self.paperDoc = None
        self.paperTemplate = '江苏科技大学试卷模板.docx'

    def convert2QuestionChapterNums(self):
        questionChapterNums = []

        for chapterQuestionNum in self.chapterQuestionNums:
            index = 0
            for questionTypeNum in chapterQuestionNum:
                if index + 1 > len(questionChapterNums):
                    questionChapterNums.append([])
                questionChapterNums[index].append(questionTypeNum)
                index = index + 1
        self.questionChpaterNums = questionChapterNums
        self.questionNums = []
        for questionChapterNum in self.questionChpaterNums:
            self.questionNums.append(sum(questionChapterNum))

    def selectQuestionNo(self):
        chapterIndex = 0
        self.chapterQuestionNames=[]
        for chapterQuestionNum in self.chapterQuestionNums:
            chapterQuestionTypeIndex = 0
            chapterQuestiontypeQuestionName = []
            for questionType in self.questionTypes:
                chapterName = self.subjectInfo.chapters[chapterIndex]
                chapterPath = os.path.join(self.subjectInfo.subjectPath, chapterName)
                chapterQuestiontypePath = os.path.join(chapterPath, questionType)
                if os.path.exists(chapterQuestiontypePath):
                    chapterQuestiontypeQuestionNameItem = os.listdir(chapterQuestiontypePath)
                    random.shuffle(chapterQuestiontypeQuestionNameItem)
                    chapterQuestiontypeQuestionNameItem = chapterQuestiontypeQuestionNameItem[
                                                          0:chapterQuestionNum[chapterQuestionTypeIndex]]
                else:
                    chapterQuestiontypeQuestionNameItem=[]
                chapterQuestiontypeQuestionName.append(chapterQuestiontypeQuestionNameItem)
                chapterQuestionTypeIndex = chapterQuestionTypeIndex + 1
            chapterIndex = chapterIndex + 1
            self.chapterQuestionNames.append(chapterQuestiontypeQuestionName)

        self.convert2QuestionTypeQuestionNames()

    def convert2QuestionTypeQuestionNames(self):
        '''
        转换成题型列表；即设置self.questionTypeQuestionNames为list[list[(chapterName,questionName)]]
        :return:
        '''
        questionTypeQuestionNames = []
        chapterIndex = 0
        for chapterQuestionName in self.chapterQuestionNames:
            questionTypeIndex = 0
            for chapterQuestionNameItem in chapterQuestionName:
                if questionTypeIndex + 1 > len(questionTypeQuestionNames):
                    questionTypeQuestionNames.append([])
                for questionIndex in range(len(chapterQuestionNameItem)):
                    chapterQuestionNameItem[questionIndex] = (
                        self.subjectInfo.chapters[chapterIndex], chapterQuestionNameItem[questionIndex])
                questionTypeQuestionNames[questionTypeIndex].extend(chapterQuestionNameItem)
                questionTypeIndex = questionTypeIndex + 1
            chapterIndex = chapterIndex + 1
        self.questionTypeQuestionNames = questionTypeQuestionNames

    def generatePaper(self):
        self.paperDoc = docx.Document(self.paperTemplate)
        # self.paperDoc.add_heading(self.subjectInfo.subjectName, 0)
        questionTypeChineseIndex = 0
        self.outputQuestionPath=[]
        for _, questionTypeIndex in enumerate(self.subjectInfo.questionTypesOrder):

            questionTypeItem=self.subjectInfo.questionTypes[questionTypeIndex]
            if self.questionNums[questionTypeIndex]==0:
                continue
            questionTypeTitleParagraph = self.paperDoc.add_paragraph()
            questionTypeTitleParagraph.text = questionChineseNo[questionTypeChineseIndex] + questionTypeItem
            questionNo = 1
            for question in self.questionTypeQuestionNames[questionTypeIndex]:
                questionPath = os.path.join(self.subjectInfo.subjectPath,  question[0], questionTypeItem,
                                            question[1])
                print(questionPath)
                self.outputQuestionPath.append(questionPath)
                sourceDoc = docx.Document(questionPath)
                self.paperDoc = utils.docAdding.docAdding(sourceDoc, self.paperDoc, questionNo)
                questionNo += 1
                self.paperDoc.add_paragraph()  # 题目之间用段落分开
            questionTypeChineseIndex = questionTypeChineseIndex + 1
        #self.paperDoc.save(self.subjectInfo.subjectName + '试卷.docx')


if __name__ == '__main__':
    subjectInfo = statistics.subject.SubjectInfo('信息论')
    subjectInfo.getSubjectStatics()

    for i in range(len(subjectInfo.chapters)):
        print(subjectInfo.chapters[i] + ':')
        chapterInfoStr=''
        for j in range(len(subjectInfo.questionTypes)):
            chapterInfoStr=chapterInfoStr+'\t' + subjectInfo.questionTypes[j] + ': ' + str(subjectInfo.chapterQuestionNums[i][j])
            #print('\t' + subjectInfo.questionTypes[j] + ': ' + str(subjectInfo.chapterQuestionNums[i][j]))
        print(chapterInfoStr)
    paperInfo = PaperInfo(subjectInfo=subjectInfo)

    paperInfo.questionTypes = ['选择题', '填空题', '判断题', '综合题']
    paperInfo.questionNums = [20, 10, 10, 5]
    paperInfo.chapterQuestionNums = [[3, 2, 0, 1], [13, 4, 3, 4], [4, 4, 7, 0]]
    paperInfo.convert2QuestionChapterNums()
    paperInfo.selectQuestionNo()
    paperInfo.generatePaper()
