import docx
import os
import utils.docAdding
import utils.utils


class QuestionMerging:
    def __init__(self):
        self.destDoc=docx.Document()

    def merge(self, subjectName, subjectPath):
        # 按照课程名，章节，题型合并多个题目到一个word文档
        # isMergeToOneFile=True：所有的题目合并为一个word文档
        # isMergeToOneFile=False：一个课程/章节/题型目录下的题目合并为一个word文档，合并后命名为：‘课程_章节_题型.docx'

        if subjectName is None:
            print('必须要有课程名称')
            return

        if subjectPath is None:
            print('必须要有课程试题库路径')
            return

        self.destDoc.add_heading(subjectName, 0)
        chapters = os.listdir(subjectPath)

        self._mergeToOneFileInSubject(subjectName, subjectPath, chapters)

        return self.destDoc




    def _mergeToOneFileInSubject(self, subjectName, subjectPath, chapters):
        for chapterItem in chapters:
            self.destDoc.add_heading(chapterItem, 1)
            chapterItemPath = os.path.join(subjectPath, chapterItem)
            questionTypes = os.listdir(chapterItemPath)

            self._mergeToOneFileInSubjectChapter(subjectName, chapterItemPath, questionTypes)
            self.destDoc.add_paragraph()  # 章节和章节之间用 空段落分割


    def _mergeToOneFileInSubjectChapter(self, subjectName, subjectChapterPath, questionTypes):
        for questionTypeItem in questionTypes:
            self.destDoc.add_heading(questionTypeItem, 2)
            chapterQuestionTypePath = os.path.join(subjectChapterPath, questionTypeItem)
            self._mergeToOneFileInSubjectCharpterQuestionType(chapterQuestionTypePath)



    def _mergeToOneFileInSubjectCharpterQuestionType(self, chapterQuestionTypePath):
        questions = os.listdir(chapterQuestionTypePath)
        questions = utils.utils.moveNonDocxFile(questions)
        questionNo = 1

        for questionItem in questions:
            questionItemPath = os.path.join(chapterQuestionTypePath, questionItem)
            sourceDoc = docx.Document(questionItemPath)
            self.destDoc = utils.docAdding.docAdding(sourceDoc, self.destDoc, questionNo)
            questionNo += 1
            self.destDoc.add_paragraph()  # 题目之间用段落分开




if __name__ == '__main__':

    questionMerging = QuestionMerging()
    mergedDoc = questionMerging.merge(['信息论'], isMergeToOneFile=True)
    # mergedDoc = questionMerging.merge(['示例课程'], isMergeToOneFile=True)
    # mergedDoc = quesMerging.merge(['信息论','示例课程'], chapters=None, questionTypes=None, isMergeToOneFile=True)
    # mergedDoc = quesMerging.merge(['示例课程'], chapters=['CH1'], questionTypes=['选择题'], isMergeToOneFile=True)
    # mergedDoc = quesMerging.merge(['信息论'], chapters=None, questionTypes=None, isMergeToOneFile=True)
    if len(mergedDoc) == 1:
        mergedDoc[0].save(mergedDoc[0].paragraphs[0].text + '.docx')
    else:
        for docItem in mergedDoc:
            docItem.save(docItem.paragraphs[0].text + '-' + docItem.paragraphs[1].text + '-' + docItem.paragraphs[
                2].text + '.docx')
