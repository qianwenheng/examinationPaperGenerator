import docx
import lxml
import zipfile
import os
import copy
import utils.updateableZipFile
import win32com
import win32com.client as win32


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = docx.oxml.OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(docx.oxml.ns.qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(docx.oxml.ns.qn('w:{}'.format(key)), str(edge_data[key]))


def addEquationParagraph(destDoc, sourceParagraph, isAddingNo=False, questionNo=None, tableNode=None):
    '''
    将含有方程的整个一个paragrpah添加到destDoc中
    如果tableNode中含有一个表格，也将表格加在该paragraph的后面
    注意table是和paragraph同等级的
    :param destDoc:
    :param sourceParagraph:
    :param isAddingNo:
    :param tableNode:
    :return:
    '''

    if os.path.exists('temp.docx'):
        os.remove('temp.docx')
    if os.path.exists('document.xml'):
        os.remove('document.xml')
    sourceParagraphXml = sourceParagraph._element.xml
    destDoc.save('temp.docx')
    z = zipfile.ZipFile('temp.docx')
    destTree = lxml.etree.parse(z.open('word/document.xml'))
    sourceParagraphTree = lxml.etree.ElementTree(
        lxml.etree.fromstring(sourceParagraphXml))  # 将方程所在段的xml转成ElementTree结构
    sourceParagraphRootNode = sourceParagraphTree.getroot()  # 将方程所在段的xml对应的ElementTree结构转成Element结构

    for node in destTree.iter():
        if 'sectPr' in node.tag:  # 当前文档的结尾
            node.addprevious(sourceParagraphRootNode)
            if tableNode is not None:
                node.addprevious(tableNode)

    destTree.write("document.xml", encoding="utf-8", xml_declaration=True, standalone="yes", pretty_print=True)

    with utils.updateableZipFile.UpdateableZipFile("temp.docx", "a") as o:
        o.write("document.xml", "word/document.xml")

    destDoc = docx.Document('temp.docx')

    if isAddingNo:
        destParagraph = destDoc.paragraphs[-1]
        destRun = destParagraph.runs[0]
        questionNoRunElement = destParagraph._element._new_r()
        destRun._element.addprevious(questionNoRunElement)
        questionNoRun = docx.text.run.Run(questionNoRunElement, destRun._parent)
        questionNoRun.text = str(questionNo) + '. '

    os.remove('document.xml')
    z.close()
    os.remove('temp.docx')
    return destDoc


def addImageParagraph(destDoc, sourceParagraph, sourceDoc, isAddingNo=False, questionNo=None, table=None):
    destParagraph = destDoc.add_paragraph()

    for runItem in sourceParagraph.runs:
        if runItem.text is not '':  # textRun
            styleList = destDoc.styles
            if runItem.font.name is None:
                runItem.font.name = 'Times New Roman'
            sourceRunFont = runItem.font.name.encode('utf-8').decode('utf-8')
            if runItem.font.name not in styleList:
                sourceCharacterStyle = destDoc.styles.add_style(sourceRunFont,
                                                                docx.enum.style.WD_STYLE_TYPE.CHARACTER)
                sourceCharacterStyle.font.name = sourceRunFont
                sourceCharacterStyle._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), sourceRunFont)
            else:
                sourceCharacterStyle = styleList[sourceRunFont]

            destRun = destParagraph.add_run(runItem.text, style=sourceCharacterStyle)

            destRun.bold = runItem.bold
            destRun.italic = runItem.italic
            destRun.underline = runItem.underline
            destRun.font.name = sourceRunFont
            destRun.font.size = runItem.font.size
            destRun.font.color.rgb = runItem.font.color.rgb
            destRun.font.superscript = runItem.font.superscript
            destRun.font.subscript = runItem.font.subscript
        else:  # imageRun
            imageNames = []
            imageIndex = 1
            destRun = destParagraph.add_run()
            runXml = runItem._element.xml
            runTree = lxml.etree.ElementTree(lxml.etree.fromstring(runXml))
            nmspcDict = next(runTree.getroot().iter()).nsmap

            namespace = {'wp': nmspcDict['wp']}
            extentNode = runTree.findall('.//wp:extent', namespace)
            extents = []
            for nodeItem in extentNode:  # 先找到图像在word中的显示尺寸，暂存
                imageWidth = nodeItem.attrib.get('cx', 0)
                imageHeight = nodeItem.attrib.get('cy', 0)
                extents.append((imageWidth, imageHeight))
            for node in runTree.iter():
                if 'blip' in node.tag:  # 在类似<a:blip r:embed="rId6">这样的标签中找出图像对应rId
                    attribKey = '{' + nmspcDict['r'] + '}embed'
                    imageRId = node.attrib.get(attribKey, None)
                    if imageRId is not None:  # 判断imageRId不是None，主要是针对<pic:blipFill>
                        imagePart = sourceDoc.part.related_parts[imageRId]
                        byteData = imagePart._blob
                        imageName = 'image' + str(imageIndex) + ".png"
                        imageNames.append(imageName)
                        with open(imageName, "wb") as fh:
                            fh.write(byteData)
                            fh.close()
                        imageIndex += 1

                    index = 0
                    for imageName in imageNames:
                        inline_shape = destRun.add_picture(imageName)
                        inline_shape.height = int(extents[index][1])
                        inline_shape.width = int(extents[index][0])
                        index = index + 1
                        os.remove(imageName)

    destParagraph.paragraph_format.alignment = sourceParagraph.paragraph_format.alignment

    if isAddingNo:
        destParagraph = destDoc.paragraphs[-1]
        destRun = destParagraph.runs[0]
        questionNoRunElement = destParagraph._element._new_r()
        destRun._element.addprevious(questionNoRunElement)
        questionNoRun = docx.text.run.Run(questionNoRunElement, destRun._parent)
        questionNoRun.text = str(questionNo) + '. '

    if table is not None:
        destParagraph._element.addnext(copy.deepcopy(table._element))
    return destDoc


def addMathTypeEquationParagraph(destDoc, sourceParagraph, sourceDoc, isAddingNo=False, questionNo=None,
                                 tableNode=None):
    '''
    希望可以直接从source文档中抽取mathtype格式的公式输入到destDoc中，
    但是没能实现这个功能
    :param destDoc:
    :param sourceParagraph:
    :param sourceDoc:
    :param isAddingNo:
    :param questionNo:
    :param tableNode:
    :return:
    '''
    # equationWmfFiles = []
    # eqautionWmfFileIndex = 1
    # oleObjectFiles = []
    # oleObjectFileIndex = 1
    #
    # sourceParagraphXml = sourceParagraph._element.xml
    # sourceParagraphTree = lxml.etree.ElementTree(lxml.etree.fromstring(sourceParagraphXml))
    # nmspcDict = next(sourceParagraphTree.getroot().iter()).nsmap
    #
    # for node in sourceParagraphTree.iter():
    #     if 'imagedata' in node.tag:  # 在类似<a:blip r:embed="rId6">这样的标签中找出图像对应rId
    #         attribKey = '{' + nmspcDict['r'] + '}id'
    #         imageRId = node.attrib.get(attribKey, None)
    #         if imageRId is not None:  # 判断imageRId不是None，主要是针对<pic:blipFill>
    #             imagePart = sourceDoc.part.related_parts[imageRId]
    #             byteData = imagePart._blob
    #             equationWmfFile = 'image' + str(eqautionWmfFileIndex) + ".wmf"
    #             equationWmfFiles.append(equationWmfFile)
    #             with open(equationWmfFile, "wb") as fh:
    #                 fh.write(byteData)
    #                 fh.close()
    #             eqautionWmfFileIndex += 1
    #     elif 'OLEObject' in node.tag:
    #         attribKey = '{' + nmspcDict['r'] + '}id'
    #         imageRId = node.attrib.get(attribKey, None)
    #         if imageRId is not None:  # 判断imageRId不是None，主要是针对<pic:blipFill>
    #             imagePart = sourceDoc.part.related_parts[imageRId]
    #             byteData = imagePart._blob
    #             oleObjectFile = 'oleObject' + str(oleObjectFileIndex) + ".bin"
    #             oleObjectFiles.append(oleObjectFile)
    #             with open(oleObjectFile, "wb") as fh:
    #                 fh.write(byteData)
    #                 fh.close()
    #             oleObjectFileIndex += 1
    #
    # if os.path.exists('temp.docx'):
    #     os.remove('temp.docx')
    # if os.path.exists('document.xml'):
    #     os.remove('document.xml')
    #
    # destDoc.save('temp.docx')
    # z = zipfile.ZipFile('temp.docx')
    # destTree = lxml.etree.parse(z.open('word/document.xml'))
    # sourceParagraphRootNode = sourceParagraphTree.getroot()  # 将方程所在段的xml对应的ElementTree结构转成Element结构
    #
    # for node in destTree.iter():
    #     if 'sectPr' in node.tag:  # 当前文档的结尾
    #         node.addprevious(sourceParagraphRootNode)
    #         if tableNode is not None:
    #             node.addprevious(tableNode)
    #
    # destTree.write("document.xml", encoding="utf-8", xml_declaration=True, standalone="yes", pretty_print=True)
    #
    # with utils.updateableZipFile.UpdateableZipFile("temp.docx", "a") as o:
    #     o.write("document.xml", "word/document.xml")
    #     for i in range(len(equationWmfFiles)):
    #         o.write(equationWmfFiles[i], 'word/media/' + equationWmfFiles[i])
    #         o.write(oleObjectFiles[i], 'word/embeddings/' + oleObjectFiles[i])
    # destDoc = docx.Document('temp.docx')
    # os.remove('document.xml')
    # z.close()
    #
    # os.remove('temp.docx')

    return destDoc


def getTableElmentIndex(sourceDoc):
    sourceXml = sourceDoc.element.xml
    sourceTree = lxml.etree.ElementTree(lxml.etree.fromstring(sourceXml))
    bodyNode = sourceTree.getroot().getchildren()[0]

    nmspcDict = next(sourceTree.getroot().iter()).nsmap
    tableTag = '{' + nmspcDict['w'] + '}tbl'
    tableIndice = []

    nodeIndex = 0
    tableNodes = []
    for nodeItem in bodyNode.getchildren():
        if nodeItem.tag == tableTag:
            tableIndice.append(nodeIndex)
            tableNodes.append(nodeItem)
        nodeIndex = nodeIndex + 1

    return tableIndice, tableNodes


def docAdding(sourceDoc, destDoc, questionNo=None):
    isAddingNo = True  # 因为题号只会加在第一句，这个flag用来控制是否加题号

    isHasTable = False
    if len(sourceDoc.tables) > 0:
        tableIndice, tableNodes = getTableElmentIndex(sourceDoc)
        isHasTable = True

    paragraphIndex = 0

    for paragraphItem in sourceDoc.paragraphs:
        isAddTable = False
        if "m:oMathPara" in paragraphItem._element.xml or "m:oMath" in paragraphItem._element.xml:  ## 含有【显示】格式的公式的段的处理
            if isHasTable:  ##处理表格
                isAddTable = False
                for i in range(len(tableIndice)):
                    if paragraphIndex == tableIndice[i] - 1 - i:
                        destDoc = addEquationParagraph(destDoc, paragraphItem,
                                                       isAddingNo=isAddingNo,
                                                       questionNo=questionNo,
                                                       tableNode=tableNodes[i])

                        for row in destDoc.tables[-1].rows:  ##设置边框
                            for cell in row.cells:
                                set_cell_border(
                                    cell,
                                    top={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                    left={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                    right={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                )

                        isAddTable = True
                if not isAddTable:
                    destDoc = addEquationParagraph(destDoc, paragraphItem, isAddingNo=isAddingNo, questionNo=questionNo)
            else:
                destDoc = addEquationParagraph(destDoc, paragraphItem, isAddingNo=isAddingNo, questionNo=questionNo)
            isAddingNo = False  # 无论当前paragraph是否要添加行号，后面都不会再添加了

        elif "pic:pic" in paragraphItem._element.xml:
            if isHasTable:
                isAddTable = False
                for i in range(len(tableIndice)):
                    if paragraphIndex == tableIndice[i] - 1 - i:
                        table = sourceDoc.tables[i]
                        destDoc = addImageParagraph(destDoc, paragraphItem, sourceDoc, isAddingNo=isAddingNo,
                                                    questionNo=questionNo, table=table)

                        for row in destDoc.tables[-1].rows:  # 设置边框
                            for cell in row.cells:
                                set_cell_border(
                                    cell,
                                    top={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                    left={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                    right={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                )

                        isAddTable = True
                if not isAddTable:
                    destDoc = addImageParagraph(destDoc, paragraphItem, sourceDoc, isAddingNo=isAddingNo,
                                                questionNo=questionNo)
            else:
                destDoc = addImageParagraph(destDoc, paragraphItem, sourceDoc, isAddingNo=isAddingNo,
                                            questionNo=questionNo)
            isAddingNo = False

        elif 'v:formulas' in paragraphItem._element.xml:
            destDoc = addMathTypeEquationParagraph(destDoc, paragraphItem, sourceDoc, isAddingNo=isAddingNo,
                                                   questionNo=questionNo)

        else:  # 正常情况处理
            destParagraph = destDoc.add_paragraph()
            for runItem in paragraphItem.runs:
                if runItem.font.name is None:
                    runItem.font.name = 'Times New Roman'
                sourceRunFontName = runItem.font.name.encode('utf-8').decode('utf-8')
                if runItem.font.name not in destDoc.styles:
                    sourceCharacterStyle = destDoc.styles.add_style(sourceRunFontName,
                                                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER)
                    sourceCharacterStyle.font.name = sourceRunFontName
                    sourceCharacterStyle._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), sourceRunFontName)
                else:
                    sourceCharacterStyle = destDoc.styles[sourceRunFontName]

                if isAddingNo:
                    destParagraph.add_run(str(questionNo) + '. ', style=sourceCharacterStyle)
                    # destParagraph.add_run(str(questionNo) + '. ')
                    isAddingNo = False

                destRun = destParagraph.add_run(runItem.text, style=sourceCharacterStyle)
                # destRun = destParagraph.add_run(runItem.text)
                destRun.bold = runItem.bold
                destRun.italic = runItem.italic
                destRun.underline = runItem.underline
                destRun.font.name = sourceRunFontName
                destRun.font.size = runItem.font.size
                destRun.font.color.rgb = runItem.font.color.rgb
                destRun.font.superscript = runItem.font.superscript
                destRun.font.subscript = runItem.font.subscript
            destParagraph.paragraph_format.alignment = paragraphItem.paragraph_format.alignment

            if isHasTable:  ##处理表格
                for i in range(len(tableIndice)):
                    if paragraphIndex == tableIndice[i] - 1 - i:
                        table = sourceDoc.tables[i]
                        destParagraph._element.addnext(copy.deepcopy(table._element))

                        for row in destDoc.tables[-1].rows:
                            for cell in row.cells:
                                set_cell_border(
                                    cell,
                                    top={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                    left={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                    right={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                )

        paragraphIndex = paragraphIndex + 1

    return destDoc
